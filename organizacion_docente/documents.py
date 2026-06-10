from io import BytesIO
import re
from decimal import Decimal

from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from django.contrib.staticfiles import finders


from datetime import timedelta
from docx.shared import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def limpiar_nombre_archivo(texto):
    """
    Limpia un texto para usarlo como nombre de archivo.
    """
    texto = texto or "documento"
    texto = texto.strip().lower()
    texto = re.sub(r"[^a-zA-Z0-9áéíóúÁÉÍÓÚñÑ_-]+", "_", texto)
    texto = texto.strip("_")
    return texto[:80]


MESES_ES = {
    1: "enero",
    2: "febrero",
    3: "marzo",
    4: "abril",
    5: "mayo",
    6: "junio",
    7: "julio",
    8: "agosto",
    9: "septiembre",
    10: "octubre",
    11: "noviembre",
    12: "diciembre",
}


def fecha_larga_es(fecha):
    if not fecha:
        return "fecha pendiente de definir"

    return f"{fecha.day} de {MESES_ES.get(fecha.month, '')} de {fecha.year}"


def obtener_apellido_docente(nombre):
    partes = (nombre or "").strip().split()

    if not partes:
        return ""

    return partes[-1]


def texto_periodo_academico(semestre, anio):
    if semestre == "I":
        return f"primer semestre de {anio}"

    if semestre == "II":
        return f"segundo semestre de {anio}"

    if semestre == "VERANO":
        return f"verano de {anio}"

    if semestre == "ESPECIAL":
        return f"periodo especial de {anio}"

    return f"{semestre} de {anio}"


def numero_nota_sipe(organizacion):
    """
    Genera número de nota con formato:
    056-SIPE-UTPCH-2026
    """

    numero = organizacion.numero_pago or organizacion.pk or 0

    try:
        numero = int(str(numero).strip())
        numero_formateado = f"{numero:03d}"
    except ValueError:
        numero_formateado = str(numero).strip()

    return f"{numero_formateado}-SIPE-UTPCH-{organizacion.anio}"


def agregar_parrafo_nota(documento, partes, align=None, space_after=8):
    """
    partes = [
        ("texto normal", False),
        ("texto en negrita", True),
    ]
    """

    parrafo = documento.add_paragraph()

    for texto, bold in partes:
        run = parrafo.add_run(texto)
        run.bold = bold
        run.font.name = "Arial Narrow"
        run.font.size = Pt(11)

    if align:
        parrafo.alignment = align

    parrafo.paragraph_format.space_after = Pt(space_after)

    return parrafo


def agregar_bullet_nota(documento, texto, nivel=0):
    parrafo = documento.add_paragraph(style="List Bullet")
    parrafo.paragraph_format.left_indent = Inches(0.3 + (nivel * 0.25))

    run = parrafo.add_run(texto)
    run.font.name = "Arial Narrow"
    run.font.size = Pt(10)

    return parrafo


def agregar_encabezado_utp_docx(documento):
    """
    Agrega el encabezado institucional como imagen.
    Si la imagen no existe, coloca texto básico.
    """

    encabezado = finders.find("img/encabezado_utp_chiriqui.png")

    if encabezado:
        parrafo = documento.add_paragraph()
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = parrafo.add_run()
        run.add_picture(encabezado, width=Inches(7.1))
        parrafo.paragraph_format.space_after = Pt(10)
    else:
        agregar_parrafo_nota(
            documento,
            [("UNIVERSIDAD TECNOLÓGICA DE PANAMÁ", True)],
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=2,
        )

        agregar_parrafo_nota(
            documento,
            [("Centro Regional de Chiriquí", True)],
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=12,
        )



def contexto_nota_docente(organizacion):
    """
    Prepara el contexto de la nota al docente con el formato oficial.
    """

    fecha_actual = timezone.localtime(timezone.now())

    nombre_docente = organizacion.docente.nombre_completo
    cedula_docente = organizacion.cedula_docente or organizacion.docente.cedula

    codigo_asignatura = (
        organizacion.codigo_asignatura
        or organizacion.asignatura.codigo_asignatura
        or ""
    )

    codigo_horario = (
        organizacion.codigo_horario
        or organizacion.asignatura.codigo_horario
        or ""
    )

    periodo_texto = texto_periodo_academico(
        organizacion.semestre,
        organizacion.anio,
    )

    fecha_listado = organizacion.fecha_matricula

    return {
        "organizacion": organizacion,
        "fecha_actual": fecha_actual,
        "fecha_actual_texto": fecha_larga_es(fecha_actual),

        "numero_nota": numero_nota_sipe(organizacion),

        "docente": organizacion.docente,
        "nombre_docente": nombre_docente,
        "apellido_docente": obtener_apellido_docente(nombre_docente),
        "cedula_docente": cedula_docente,

        "facultad": organizacion.facultad,
        "programa": organizacion.programa,
        "asignatura": organizacion.asignatura,

        "numero_pago": organizacion.numero_pago or "",
        "anio": organizacion.anio,
        "semestre": organizacion.get_semestre_display(),
        "periodo_texto": periodo_texto,
        "grupo_aula": organizacion.grupo_aula,

        "codigo_asignatura": codigo_asignatura,
        "codigo_horario": codigo_horario,
        "total_horas": organizacion.total_horas,
        "total_creditos": organizacion.total_creditos,

        "fechas_clases": organizacion.fechas_clases or "fechas pendientes de definir",
        "horario": organizacion.horario or "horario pendiente de definir",
        "modalidad": "a distancia",

        "fecha_listado_oficial": fecha_larga_es(fecha_listado),
        "observaciones": organizacion.observaciones or "",

        "directora_centro": "Iveth Moreno",
        "correo_postgrado": "postgrados.chiriqui@utp.ac.pa",
        "coordinador_postgrado": "Mgtr. José Mendoza",
        "iniciales": "JM/elc",

        "usar_encabezado_imagen": bool(
            finders.find("img/encabezado_utp_chiriqui.png")
        ),
    }


def nombre_archivo_nota(organizacion, extension):
    docente = limpiar_nombre_archivo(organizacion.docente.nombre_completo)
    asignatura = limpiar_nombre_archivo(organizacion.asignatura.nombre)
    fecha = timezone.localtime(timezone.now()).strftime("%Y%m%d_%H%M")

    return f"nota_docente_{docente}_{asignatura}_{fecha}.{extension}"


def generar_pdf_nota_docente(organizacion, request=None):
    """
    Genera la nota al docente en PDF usando un template HTML.
    """

    from weasyprint import HTML

    context = contexto_nota_docente(organizacion)

    html_string = render_to_string(
        "organizacion_docente/documentos/nota_docente_pdf.html",
        context,
    )

    base_url = None

    if request:
        base_url = request.build_absolute_uri("/")

    pdf_file = HTML(
        string=html_string,
        base_url=base_url,
    ).write_pdf()

    filename = nombre_archivo_nota(organizacion, "pdf")

    response = HttpResponse(
        pdf_file,
        content_type="application/pdf",
    )

    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    return response


def agregar_parrafo(documento, texto="", bold=False, align=None, space_after=6):
    """
    Agrega un párrafo con formato básico a un documento Word.
    """

    parrafo = documento.add_paragraph()
    run = parrafo.add_run(texto)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(11)

    if align:
        parrafo.alignment = align

    parrafo.paragraph_format.space_after = Pt(space_after)

    return parrafo


def agregar_linea_dato(documento, etiqueta, valor):
    """
    Agrega una línea de dato en formato: Etiqueta: Valor
    """

    parrafo = documento.add_paragraph()
    parrafo.paragraph_format.space_after = Pt(3)

    run_label = parrafo.add_run(f"{etiqueta}: ")
    run_label.bold = True
    run_label.font.name = "Arial"
    run_label.font.size = Pt(10)

    run_value = parrafo.add_run(str(valor or "-"))
    run_value.font.name = "Arial"
    run_value.font.size = Pt(10)

    return parrafo


def generar_docx_nota_docente(organizacion):
    """
    Genera la nota al docente en Word DOCX con el formato oficial:
    1. Nota principal
    2. Carta compromiso
    3. Observaciones
    """

    context = contexto_nota_docente(organizacion)

    documento = Document()

    section = documento.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    # ========================================================
    # PÁGINA 1 — NOTA PRINCIPAL
    # ========================================================

    agregar_encabezado_utp_docx(documento)

    agregar_parrafo_nota(
        documento,
        [(context["numero_nota"], True)],
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=10,
    )

    agregar_parrafo_nota(
        documento,
        [(f"David, {context['fecha_actual_texto']}", False)],
        space_after=28,
    )

    agregar_parrafo_nota(documento, [("Magister", False)], space_after=1)
    agregar_parrafo_nota(documento, [(context["nombre_docente"], False)], space_after=1)
    agregar_parrafo_nota(documento, [("Docente de Maestría", False)], space_after=1)
    agregar_parrafo_nota(documento, [("E. S. M.", False)], space_after=22)

    agregar_parrafo_nota(
        documento,
        [(f"Respetado magister {context['apellido_docente']}:", False)],
        space_after=18,
    )

    agregar_parrafo_nota(
        documento,
        [
            (
                "La Universidad Tecnológica de Panamá, Centro Regional de Chiriquí, "
                "se complace en contar con sus servicios como docente en la ",
                False,
            ),
            (context["programa"].nombre, True),
            (f", correspondiente al {context['periodo_texto']}.", False),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )

    agregar_parrafo_nota(
        documento,
        [
            ("En esta ocasión, usted tendrá asignado el curso ", False),
            (f"{context['asignatura'].nombre} ({context['codigo_asignatura']})", True),
            (
                f", el cual se impartirá en modalidad {context['modalidad']} "
                f"en el horario {context['horario']}, {context['fechas_clases']}. ",
                False,
            ),
            (
                "El listado oficial estará disponible en su perfil de docente "
                "(matricula.utp.ac.pa) a partir del ",
                False,
            ),
            (f"{context['fecha_listado_oficial']}.", True),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )

    agregar_parrafo_nota(
        documento,
        [
            (
                "Adicionalmente, le informamos que, de acuerdo con los lineamientos "
                "de la VIPE, es necesario que los docentes de postgrado firmen una "
                "carta compromiso que confirme la disponibilidad en los horarios y "
                "fechas acordadas. Adjuntamos el formato de la carta, la cual deberá "
                "firmar a manuscrito y remitir al correo ",
                False,
            ),
            (context["correo_postgrado"], False),
            (
                " o entregarla en la Coordinación de Postgrado del Centro Regional "
                "de Chiriquí.",
                False,
            ),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=12,
    )

    agregar_parrafo_nota(
        documento,
        [
            (
                "Agradecemos su disposición para compartir sus conocimientos y le "
                "deseamos un exitoso desarrollo del curso. No dude en contactarnos "
                "para cualquier consulta o apoyo adicional que requiera.",
                False,
            )
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=22,
    )

    agregar_parrafo_nota(documento, [("Atentamente,", False)], space_after=38)

    agregar_parrafo_nota(
        documento,
        [(context["coordinador_postgrado"], True)],
        space_after=1,
    )

    agregar_parrafo_nota(
        documento,
        [("Coordinador de Postgrado", False)],
        space_after=1,
    )

    agregar_parrafo_nota(
        documento,
        [("Centro Regional de Chiriquí", False)],
        space_after=16,
    )

    agregar_parrafo_nota(documento, [(context["iniciales"], False)], space_after=8)

    agregar_parrafo_nota(documento, [("Adj.: Carta compromiso", False)], space_after=1)
    agregar_parrafo_nota(documento, [("      Lista Preliminar de estudiantes", False)], space_after=1)
    agregar_parrafo_nota(documento, [("      Observaciones", False)], space_after=1)

    documento.add_page_break()

    # ========================================================
    # PÁGINA 2 — CARTA COMPROMISO
    # ========================================================

    agregar_parrafo_nota(
        documento,
        [("Panamá, ____ de _______________ de __________", False)],
        space_after=70,
    )

    agregar_parrafo_nota(documento, [("Doctora", False)], space_after=1)
    agregar_parrafo_nota(documento, [(context["directora_centro"], False)], space_after=1)
    agregar_parrafo_nota(documento, [("Directora", False)], space_after=1)
    agregar_parrafo_nota(documento, [("Centro Regional de Chiriquí", False)], space_after=1)
    agregar_parrafo_nota(documento, [("Universidad Tecnológica de Panamá", False)], space_after=28)

    agregar_parrafo_nota(
        documento,
        [("Respetada Dra. Moreno:", False)],
        space_after=18,
    )

    agregar_parrafo_nota(
        documento,
        [
            ("Por este medio me comprometo a dictar el curso ", False),
            (f"{context['asignatura'].nombre} ({context['codigo_asignatura']})", True),
            (", del programa ", False),
            (context["programa"].nombre, True),
            (f", en el {context['periodo_texto']}, ofertado por la ", False),
            (str(context["facultad"]), False),
            (
                f", que se desarrollará {context['fechas_clases']}, "
                f"en un horario de {context['horario']}. ",
                False,
            ),
            (
                "A la vez, hago constar que no tengo compromisos laborales adquiridos "
                "con ninguna institución pública, ni privada, en las fechas y horarios "
                "antes descritos.",
                False,
            ),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=38,
    )

    agregar_parrafo_nota(documento, [("Atentamente,", False)], space_after=58)

    agregar_parrafo_nota(
        documento,
        [("__________________________________", False)],
        space_after=18,
    )

    agregar_parrafo_nota(
        documento,
        [("Nombre del docente: _________________________", False)],
        space_after=18,
    )

    agregar_parrafo_nota(
        documento,
        [("Cédula: _____________________", False)],
        space_after=8,
    )

    documento.add_page_break()

    # ========================================================
    # PÁGINA 3 — OBSERVACIONES
    # ========================================================

    agregar_parrafo_nota(
        documento,
        [("OBSERVACIONES", True)],
        space_after=2,
    )

    agregar_parrafo_nota(
        documento,
        [("1. Documentos requeridos para el trámite de organización docente:", True)],
        space_after=4,
    )

    agregar_bullet_nota(documento, "Carta Compromiso firmada.")
    agregar_bullet_nota(documento, "Cédula actualizada.")


    agregar_bullet_nota(
        documento,
        "Certificación de horario laboral, si labora en otras instituciones públicas, "
        "firmada por una autoridad o jefe inmediato.",
    )

    agregar_parrafo_nota(
        documento,
        [("2. Observaciones Adicionales:", True)],
        space_after=4,
    )

    agregar_bullet_nota(
        documento,
        "El docente dispone de 7 días hábiles, a partir de la última sesión, "
        "para registrar la calificación.",
    )

    agregar_bullet_nota(
        documento,
        f"Enviar al correo {context['correo_postgrado']}, en formato PDF, "
        "un portafolio que contenga lo siguiente:",
    )

    agregar_bullet_nota(
        documento,
        "Portada con las generales de la asignatura, el docente y el grupo correspondiente.",
        nivel=1,
    )
    agregar_bullet_nota(documento, "Horario de clases.", nivel=1)
    agregar_bullet_nota(documento, "Programación del curso.", nivel=1)
    agregar_bullet_nota(
        documento,
        "Principales actividades realizadas: tareas, asignaciones, trabajos en clases, "
        "trabajos en línea, talleres, investigaciones, laboratorio, prácticas y trabajos "
        "complementarios.",
        nivel=1,
    )
    agregar_bullet_nota(
        documento,
        "Ejercicios cortos, quiz y parciales o pruebas formativas.",
        nivel=1,
    )
    agregar_bullet_nota(documento, "Examen o proyecto final.", nivel=1)
    agregar_bullet_nota(documento, "Registro de calificaciones.", nivel=1)
    agregar_bullet_nota(
        documento,
        "Anexo: proyectos finales, investigaciones, reportes o artículos realizados "
        "por los estudiantes.",
        nivel=1,
    )

    agregar_bullet_nota(
        documento,
        (
            "El portafolio debe enviarse rotulado con los siguientes datos en el nombre "
            "del archivo: nombre.apellido."
            f"{context['semestre'].upper().replace(' ', '-')}-{context['anio']}-CRCH."
        ),
    )

    output = BytesIO()
    documento.save(output)
    output.seek(0)

    filename = nombre_archivo_nota(organizacion, "docx")

    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    return response

# ============================================================
# Calendario de matrícula y pago para estudiantes
# ============================================================

def formatear_fecha_calendario(fecha):
    if not fecha:
        return "Por definir"

    return f"{fecha.day} de {MESES_ES.get(fecha.month, '')} de {fecha.year}"


def formatear_rango_fechas(inicio, fin):
    if not inicio or not fin:
        return "Por definir"

    if inicio.month == fin.month and inicio.year == fin.year:
        return (
            f"Del {inicio.day} al {fin.day} de "
            f"{MESES_ES.get(inicio.month, '')} de {inicio.year}"
        )

    return (
        f"Del {inicio.day} de {MESES_ES.get(inicio.month, '')} de {inicio.year} "
        f"al {fin.day} de {MESES_ES.get(fin.month, '')} de {fin.year}"
    )


def calcular_fechas_calendario_estudiantes(organizacion):
    """
    Obtiene las fechas del calendario de matrícula y pago.

    Prioridad:
    1. Usa los campos de texto escritos manualmente.
    2. Si están vacíos, calcula fechas automáticas desde fecha_matricula.
    """

    fecha_base = organizacion.fecha_matricula

    fechas_auto = {
        "matricula": "Por definir",
        "primer_pago": "Por definir",
        "segundo_pago": "Por definir",
        "tercer_pago": "Por definir",
        "retiro_inclusion": "Por definir",
        "retiro_fuera": "Por definir",
    }

    if fecha_base:
        fechas_auto = {
            "matricula": formatear_fecha_calendario(fecha_base),
            "primer_pago": formatear_rango_fechas(
                fecha_base,
                fecha_base + timedelta(days=13),
            ),
            "segundo_pago": formatear_rango_fechas(
                fecha_base + timedelta(days=14),
                fecha_base + timedelta(days=28),
            ),
            "tercer_pago": formatear_rango_fechas(
                fecha_base + timedelta(days=29),
                fecha_base + timedelta(days=40),
            ),
            "retiro_inclusion": formatear_rango_fechas(
                fecha_base + timedelta(days=5),
                fecha_base + timedelta(days=11),
            ),
            "retiro_fuera": formatear_rango_fechas(
                fecha_base + timedelta(days=12),
                fecha_base + timedelta(days=38),
            ),
        }

    return {
        "matricula": organizacion.fecha_matricula_texto or fechas_auto["matricula"],
        "primer_pago": organizacion.primer_pago_texto or fechas_auto["primer_pago"],
        "segundo_pago": organizacion.segundo_pago_texto or fechas_auto["segundo_pago"],
        "tercer_pago": organizacion.tercer_pago_texto or fechas_auto["tercer_pago"],
        "retiro_inclusion": organizacion.retiro_inclusion_texto or fechas_auto["retiro_inclusion"],
        "retiro_fuera": organizacion.retiro_fuera_texto or fechas_auto["retiro_fuera"],
    }


def contexto_calendario_pago(organizacion, cuotas=1):
    """
    Contexto para el calendario de matrícula y pago de estudiantes.
    Se mantiene el nombre de la función para no cambiar las vistas.
    """

    fechas_pago = calcular_fechas_calendario_estudiantes(organizacion)

    periodo_texto = texto_periodo_academico(
        organizacion.semestre,
        organizacion.anio,
    ).upper()

    codigo_asignatura = (
        organizacion.codigo_asignatura
        or organizacion.asignatura.codigo_asignatura
        or ""
    )

    codigo_horario = (
        organizacion.codigo_horario
        or organizacion.asignatura.codigo_horario
        or ""
    )

    return {
        "organizacion": organizacion,
        "fecha_actual": timezone.localtime(timezone.now()),

        "programa": organizacion.programa,
        "programa_nombre": organizacion.programa.nombre.upper(),
        "grupo_aula": organizacion.grupo_aula,
        "periodo_texto": periodo_texto,

        "asignatura": organizacion.asignatura,
        "asignatura_nombre": organizacion.asignatura.nombre,
        "codigo_asignatura": codigo_asignatura,
        "codigo_horario": codigo_horario,
        "creditos": organizacion.total_creditos,
        "docente": organizacion.docente,
        "docente_nombre": organizacion.docente.nombre_completo,

        "fechas_clases": organizacion.fechas_clases or "Por definir",
        "horario": organizacion.horario or "Por definir",

        "matricula": fechas_pago["matricula"],
        "primer_pago": fechas_pago["primer_pago"],
        "segundo_pago": fechas_pago["segundo_pago"],
        "tercer_pago": fechas_pago["tercer_pago"],
        "retiro_inclusion": fechas_pago["retiro_inclusion"],
        "retiro_fuera": fechas_pago["retiro_fuera"],

        "observacion_recargo": (
            "Recargo del 25% por pago tardío al vencimiento de cada tercio."
        ),

        "url_pago_utp": (
            "https://serviciosonline.utp.ac.pa/Pagos/UTP/bienvenida/2021/fJC8KWm8WkyUaoVnnU3R1g"
        ),
    }


def nombre_archivo_calendario(organizacion, extension):
    programa = limpiar_nombre_archivo(organizacion.programa.nombre)
    asignatura = limpiar_nombre_archivo(organizacion.asignatura.nombre)
    fecha = timezone.localtime(timezone.now()).strftime("%Y%m%d_%H%M")

    return f"calendario_matricula_pago_{programa}_{asignatura}_{fecha}.{extension}"


def set_cell_shading(cell, fill):
    """
    Aplica color de fondo a una celda de Word.
    fill debe ser HEX sin #. Ejemplo: FFFF00
    """

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=9):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    paragraph = cell.paragraphs[0]
    paragraph.alignment = align

    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Arial Narrow"
    run.font.size = Pt(size)


def aplicar_bordes_tabla_word(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()

            tc_borders = tc_pr.first_child_found_in("w:tcBorders")

            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)

            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = tc_borders.find(qn(tag))

                if element is None:
                    element = OxmlElement(tag)
                    tc_borders.append(element)

                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "6")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "000000")


def generar_pdf_calendario_pago(organizacion, request=None, cuotas=1):
    """
    Genera el calendario de matrícula y pago de estudiantes en PDF.
    """

    from weasyprint import HTML

    context = contexto_calendario_pago(
        organizacion=organizacion,
        cuotas=cuotas,
    )

    html_string = render_to_string(
        "organizacion_docente/documentos/calendario_pago_pdf.html",
        context,
    )

    base_url = request.build_absolute_uri("/") if request else None

    pdf_file = HTML(
        string=html_string,
        base_url=base_url,
    ).write_pdf()

    filename = nombre_archivo_calendario(organizacion, "pdf")

    response = HttpResponse(
        pdf_file,
        content_type="application/pdf",
    )

    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    return response


def generar_docx_calendario_pago(organizacion, cuotas=1):
    """
    Genera el calendario de matrícula y pago de estudiantes en Word DOCX.
    """

    context = contexto_calendario_pago(
        organizacion=organizacion,
        cuotas=cuotas,
    )

    documento = Document()

    section = documento.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    # Encabezado
    agregar_parrafo_nota(
        documento,
        [("UNIVERSIDAD TECNOLÓGICA DE PANAMÁ", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )

    agregar_parrafo_nota(
        documento,
        [("CENTRO REGIONAL DE CHIRIQUÍ", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )

    agregar_parrafo_nota(
        documento,
        [("SUBDIRECCIÓN DE INVESTIGACIÓN, POSTGRADO Y EXTENSIÓN", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
    )

    agregar_parrafo_nota(
        documento,
        [("COORDINACIÓN DE POSTGRADO/MAESTRÍA", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=22,
    )

    agregar_parrafo_nota(
        documento,
        [("PROGRAMA: ", True), (context["programa_nombre"], True)],
        space_after=0,
    )

    agregar_parrafo_nota(
        documento,
        [("GRUPO: ", True), (context["grupo_aula"], True)],
        space_after=0,
    )

    agregar_parrafo_nota(
        documento,
        [("PERIODO: ", True), (context["periodo_texto"], True)],
        space_after=18,
    )

    agregar_parrafo_nota(
        documento,
        [("PROGRAMACIÓN", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # Tabla programación
    tabla_prog = documento.add_table(rows=2, cols=7)
    tabla_prog.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla_word(tabla_prog)

    headers = [
        "ASIGNATURA",
        "COD.\nASIG.",
        "COD.\nHOR",
        "Cr.",
        "DOCENTE",
        "FECHAS DE CLASE",
        "HORARIO",
    ]

    for i, header in enumerate(headers):
        set_cell_text(
            tabla_prog.rows[0].cells[i],
            header,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            size=8,
        )

    valores = [
        context["asignatura_nombre"],
        context["codigo_asignatura"],
        context["codigo_horario"],
        context["creditos"],
        context["docente_nombre"],
        context["fechas_clases"],
        context["horario"],
    ]

    for i, valor in enumerate(valores):
        align = WD_ALIGN_PARAGRAPH.CENTER if i in [1, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(
            tabla_prog.rows[1].cells[i],
            valor,
            bold=False,
            align=align,
            size=8,
        )

    agregar_parrafo_nota(
        documento,
        [("CALENDARIO DE MATRÍCULA Y PAGO", True)],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    # Tabla calendario
    tabla_pago = documento.add_table(rows=8, cols=2)
    tabla_pago.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla_word(tabla_pago)

    # Header
    set_cell_shading(tabla_pago.rows[0].cells[0], "D9D9D9")
    set_cell_shading(tabla_pago.rows[0].cells[1], "D9D9D9")
    set_cell_text(tabla_pago.rows[0].cells[0], "DETALLE", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(tabla_pago.rows[0].cells[1], "FECHA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Matrícula
    set_cell_shading(tabla_pago.rows[1].cells[0], "D9D9D9")
    set_cell_shading(tabla_pago.rows[1].cells[1], "D9D9D9")
    set_cell_text(tabla_pago.rows[1].cells[0], "MATRÍCULA", bold=True)
    set_cell_text(tabla_pago.rows[1].cells[1], context["matricula"], bold=True)

    # Periodos
    merged = tabla_pago.rows[2].cells[0].merge(tabla_pago.rows[2].cells[1])
    set_cell_shading(merged, "FFFF00")
    set_cell_text(
        merged,
        "PERÍODOS DE PAGO DE MATRÍCULA",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )

    set_cell_text(
        tabla_pago.rows[3].cells[0],
        "• Pago total con descuento del 5% o Primer\n  Tercio de la Matrícula",
        bold=False,
    )
    set_cell_text(tabla_pago.rows[3].cells[1], context["primer_pago"])

    set_cell_text(
        tabla_pago.rows[4].cells[0],
        "• Pago del Segundo tercio de la matrícula",
    )
    set_cell_text(tabla_pago.rows[4].cells[1], context["segundo_pago"])

    set_cell_text(
        tabla_pago.rows[5].cells[0],
        "• Pago del Tercer Tercio de la matrícula",
    )
    set_cell_text(tabla_pago.rows[5].cells[1], context["tercer_pago"])

    # Retiro / inclusión
    set_cell_shading(tabla_pago.rows[6].cells[0], "FFFF00")
    set_cell_shading(tabla_pago.rows[6].cells[1], "FFFF00")
    set_cell_text(
        tabla_pago.rows[6].cells[0],
        "Retiro/Inclusión de asignaturas\n"
        "(Devolución del 50% del costo total de\n"
        "matrícula por retiro en este periodo)",
        bold=True,
    )
    set_cell_text(tabla_pago.rows[6].cells[1], context["retiro_inclusion"], bold=True)

    # Retiro fuera
    set_cell_shading(tabla_pago.rows[7].cells[0], "FFFF00")
    set_cell_shading(tabla_pago.rows[7].cells[1], "FFFF00")
    set_cell_text(
        tabla_pago.rows[7].cells[0],
        "Retiro fuera del periodo\n(Sin devolución)",
        bold=True,
    )
    set_cell_text(tabla_pago.rows[7].cells[1], context["retiro_fuera"], bold=True)

    # Observación como tabla aparte
    tabla_obs = documento.add_table(rows=1, cols=2)
    tabla_obs.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla_word(tabla_obs)

    set_cell_text(tabla_obs.rows[0].cells[0], "OBSERVACIÓN", bold=True)
    set_cell_text(tabla_obs.rows[0].cells[1], context["observacion_recargo"])

    agregar_parrafo_nota(
        documento,
        [
            ("Nota: ", True),
            (
                "Les exhortamos a realizar la matrícula y los pagos de cada tercio, "
                "de acuerdo con las fechas indicadas en este calendario, para evitar recargos.",
                False,
            ),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=16,
    )

    agregar_parrafo_nota(
        documento,
        [("Modalidades para pago de matrícula:", True)],
        space_after=4,
    )

    agregar_parrafo_nota(
        documento,
        [
            ("1. Pago por Tarjeta de débito y crédito desde el sitio web de la UTP.\n", False),
            (context["url_pago_utp"], False),
        ],
        space_after=8,
    )

    agregar_parrafo_nota(
        documento,
        [
            ("2. Caja UTP: ", True),
            ("Deben ser días hábiles. ", True),
            (
                "Llevar impresa su constancia de matrícula y de preferencia el calendario de pago.\n",
                False,
            ),
            (
                "Horario de atención de la caja de la UTP CR. de Chiriquí: "
                "lunes a viernes: 8:30 a.m. a 3:30 p.m.",
                True,
            ),
        ],
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=0,
    )

    output = BytesIO()
    documento.save(output)
    output.seek(0)

    filename = nombre_archivo_calendario(organizacion, "docx")

    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    return response